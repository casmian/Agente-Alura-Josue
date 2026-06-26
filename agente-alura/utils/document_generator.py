import os
import sys
import json
import argparse
import re

# Evitar errores de codificación en consolas Windows (CP1252)
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

# Permitir importaciones relativas
sys.path.append(os.path.join(os.path.dirname(__file__), ".."))

try:
    import docx
except ImportError:
    print("[WARNING] La librería 'python-docx' no está instalada o no se encuentra en el entorno virtual.")
    print("Por favor, asegúrate de activar el entorno virtual (.venv) e instalar los requerimientos.")

def clean_filename(name: str) -> str:
    """
    Limpia un nombre de cadena para que sea un nombre de archivo seguro y siga el estándar.
    Reemplaza espacios por guiones bajos y elimina caracteres especiales.
    """
    # Remover tildes y caracteres especiales
    name = re.sub(r'[áäâà]', 'a', name, flags=re.IGNORECASE)
    name = re.sub(r'[éëêè]', 'e', name, flags=re.IGNORECASE)
    name = re.sub(r'[íïîì]', 'i', name, flags=re.IGNORECASE)
    name = re.sub(r'[óöôò]', 'o', name, flags=re.IGNORECASE)
    name = re.sub(r'[úüûù]', 'u', name, flags=re.IGNORECASE)
    name = re.sub(r'[ñ]', 'n', name, flags=re.IGNORECASE)
    
    # Dejar solo caracteres alfa-numéricos, espacios y guiones
    name = re.sub(r'[^a-zA-Z0-9\s_\-]', '', name)
    # Reemplazar espacios por guiones bajos
    name = re.sub(r'[\s\-]+', '_', name)
    return name.strip('_')

def generate_markdown(doc_data: dict, output_path: str):
    """
    Genera un archivo Markdown a partir de los datos estructurados.
    """
    lines = []
    
    # Título principal
    titulo = doc_data.get("titulo", "Documento Corporativo")
    institucion = doc_data.get("institucion", "Neouniverse")
    
    for sec in doc_data.get("secciones", []):
        nivel = sec.get("nivel", 1)
        sec_titulo = sec.get("titulo", "")
        contenido = sec.get("contenido", "")
        
        # Generar cabeceras según el nivel
        hashes = "#" * nivel
        lines.append(f"{hashes} {sec_titulo}\n")
        lines.append(f"{contenido}\n")
        lines.append("---" if nivel == 1 else "") # Separador para H1
        lines.append("")
        
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
        
    print(f"[OK] Documento Markdown generado con éxito en: {output_path}")

def generate_docx(doc_data: dict, output_path: str):
    """
    Genera un archivo Word (.docx) a partir de los datos estructurados.
    Utiliza estilos Heading para garantizar compatibilidad con el document_parser de Alura.
    """
    doc = docx.Document()
    
    # Agregar título y secciones
    for sec in doc_data.get("secciones", []):
        nivel = sec.get("nivel", 1)
        sec_titulo = sec.get("titulo", "")
        contenido = sec.get("contenido", "")
        
        # En docx, agregamos cabeceras con el estilo correspondiente
        # Heading 1, Heading 2, etc.
        style_name = f"Heading {nivel}"
        try:
            doc.add_paragraph(sec_titulo, style=style_name)
        except Exception:
            # Fallback si el estilo no existe en la plantilla por defecto de python-docx
            p = doc.add_paragraph()
            run = p.add_run(sec_titulo)
            run.bold = True
            if nivel == 1:
                run.font.size = docx.shared.Pt(18)
            else:
                run.font.size = docx.shared.Pt(14)
                
        # Agregar contenido (línea por línea para manejar saltos de carro)
        for line in contenido.split('\n'):
            doc.add_paragraph(line)
            
    doc.save(output_path)
    print(f"[OK] Documento Word (.docx) generado con éxito en: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Generador de documentos para base de conocimientos del Agente Alura")
    parser.add_argument("--template", type=str, help="Ruta al archivo JSON de plantilla", default=None)
    parser.add_argument("--format", type=str, choices=["md", "docx"], help="Formato de salida (md o docx)", default=None)
    
    args = parser.parse_args()
    
    # Si no se provee plantilla, usar la plantilla de ejemplo por defecto
    template_path = args.template
    if not template_path:
        default_template = os.path.join(os.path.dirname(__file__), "plantilla_politica.json")
        if os.path.exists(default_template):
            template_path = default_template
            print(f"[INFO] Usando plantilla por defecto: {default_template}")
        else:
            print("[ERROR] No se especificó ninguna plantilla y no se encontró la plantilla por defecto.")
            print("Uso: python document_generator.py --template ruta_a_plantilla.json")
            sys.exit(1)
            
    # Leer datos del JSON
    try:
        with open(template_path, "r", encoding="utf-8") as f:
            doc_data = json.load(f)
    except Exception as e:
        print(f"[ERROR] No se pudo leer o parsear el archivo de plantilla JSON: {e}")
        sys.exit(1)
        
    # Obtener formato
    formato = args.format or doc_data.get("formato", "md").lower()
    if formato not in ["md", "docx"]:
        print(f"[WARNING] Formato '{formato}' no reconocido. Se usará 'md' por defecto.")
        formato = "md"
        
    # Configurar rutas de salida
    kb_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "base-conocimiento"))
    os.makedirs(kb_dir, exist_ok=True)
    
    # Generar nombre del archivo siguiendo el estándar: [Nombre_Limpio]__Neouniverse.[ext]
    raw_title = doc_data.get("titulo", "Documento_Nuevo")
    clean_title = clean_filename(raw_title)
    filename = f"{clean_title}__Neouniverse.{formato}"
    output_path = os.path.join(kb_dir, filename)
    
    # Generar según el formato
    if formato == "md":
        generate_markdown(doc_data, output_path)
    elif formato == "docx":
        generate_docx(doc_data, output_path)
        
    print(f"\n[ÉXITO] Archivo creado listo para ingesta: {filename}")
    
    # Preguntar si desea ejecutar el pipeline de ingesta de inmediato
    print("\n¿Deseas ejecutar el pipeline de ingesta para que el agente aprenda este documento ahora mismo?")
    try:
        # En scripts interactivos CLI, usamos input. 
        # Dado que corremos en entornos de agente, proveemos una opción automática por si acaso no es interactivo.
        response = input("Escribe 'si' para ejecutar la ingesta (Enter para omitir): ").strip().lower()
        if response in ["si", "s", "yes", "y"]:
            print("\nIniciando ingesta...")
            from agent.ingestion_service import start_ingestion
            start_ingestion()
    except Exception as e:
        print(f"\n[INFO] Omitiendo ingesta interactiva. Puedes correrla manualmente usando:")
        print(f"python agente-alura/agent/ingestion_service.py")

if __name__ == "__main__":
    main()
